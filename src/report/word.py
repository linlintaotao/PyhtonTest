# coding= utf-8

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os


class WordReporter:

    def __init__(self, path, time):
        self._time = time
        self._path = path
        self._pic = []
        self._records = None
        self._testPower = False
        self._cepResult = None

    # 获取当前文件夹下的所有png文件
    @staticmethod
    def read_file(dirName, tag='png'):
        picture = list()
        for file in os.listdir(dirName):
            if file.endswith(tag):
                picture.append(file)
        picture.sort()
        return picture

    def setRecords(self, records, testPower):
        self._records = records
        self._testPower = testPower

    def setCepResult(self, cepResult):
        self._cepResult = cepResult

    def addPng(self, doc):
        listfile = os.listdir(self._path)
        listfile.sort()
        for fileName in listfile:
            dirName = os.path.join(self._path, fileName)
            if os.path.isdir(dirName):
                doc.add_paragraph('%s 图例' % fileName, style='Intense Quote')
                pictures = self.read_file(dirName)
                for picture in pictures:
                    doc.add_picture(dirName + '/' + picture, Inches(6), Inches(4))
                doc.add_paragraph('%s 图例' % fileName, style='Intense Quote')

    def build(self):
        doc = Document()
        doc.add_heading('日常静态测试', 0)
        doc.add_paragraph('测试报告生成时间：%s' % self._time)
        doc.add_paragraph('基站: Obs 20C')
        table = None
        if self._records is not None:
            table = doc.add_table(rows=1, cols=7, style="Light Grid Accent 1")
            table.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            table.style.font.size = Pt(10)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = '串口号'
            hdr_cells[1].text = 'NaviRate'
            hdr_cells[2].text = 'WorkMode'
            hdr_cells[3].text = 'RtkDiff'
            hdr_cells[4].text = 'HeartBeat'
            hdr_cells[5].text = '总数' if self._testPower is not True else 'lisence次数'
            hdr_cells[6].text = '固定率' if self._testPower is not True else '固定次数'

            for COM, version, fixNum, percent, naviRate, workMode, heartBeatTimes, rtkDiff in self._records:
                row_cells = table.add_row().cells
                row_cells[0].text = COM
                row_cells[1].text = naviRate
                row_cells[2].text = workMode
                row_cells[3].text = rtkDiff
                row_cells[4].text = str(heartBeatTimes)
                row_cells[5].text = fixNum
                row_cells[6].text = percent

        if self._records is not None and table is not None:
            hdr_cells = table.add_row().cells
            hdr_cells[0].text = '/'
            for i in range(len(hdr_cells)):
                if i > 1:
                    hdr_cells[1].merge(hdr_cells[i])
            hdr_cells[1].text = '设备串口对应版本'

            for COM, version, fixNum, percent, naviRate, workMode, heartBeatTimes, rtkDiff in self._records:
                row_cells = table.add_row().cells
                row_cells[0].text = COM
                for i in range(len(row_cells)):
                    if i > 1:
                        row_cells[1].merge(row_cells[i])
                row_cells[1].text = version

        if self._cepResult is not None:
            cep_cells = table.add_row().cells
            cep_cells[0].text = 'Result'
            cep_cells[1].text = 'mean'
            cep_cells[2].text = 'std'
            cep_cells[3].text = '68%'
            cep_cells[4].text = '95%'
            cep_cells[5].text = '99.7%'
            cep_cells[6].text = 'max'
            for name, cepInfo in self._cepResult:
                if cepInfo is None:
                    continue
                cep_cells = table.add_row().cells
                cep_cells[0].text = name
                i = 0
                for row in cepInfo.iteritems():
                    if i == 1:
                        cep_cells[1].text = str(f'%.3f' % row[1])
                    elif i == 2:
                        cep_cells[2].text = str(f'%.3f' % row[1])
                    elif i == 5:
                        cep_cells[3].text = str(f'%.3f' % row[1])
                    elif i == 6:
                        cep_cells[4].text = str(f'%.3f' % row[1])
                    elif i == 7:
                        cep_cells[5].text = str(f'%.3f' % row[1])
                    elif i == 8:
                        cep_cells[6].text = str(f'%.3f' % row[1])
                    i += 1

        doc.add_heading('测试结果图例', level=1)
        self.addPng(doc)
        doc.save(self._path + '/' + self._time + 'report.docx')


if __name__ == '__main__':
    records = (
        ('3', '101', 'Spam', 'g'),
        ('7', '422', 'Eggs', 'h'),
        ('4', '631', 'Spam, spam, eggs, and spam', 't')
    )
    word = WordReporter(os.path.abspath('../..') + "/data", "12345678")
    word.setRecords(records, testPower=False)
    word.build()
