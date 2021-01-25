from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os


class PPReporter:

    def __init__(self, ppPath, time, savePath=None):
        self.ppPath = ppPath
        self.time = time
        self.savePath = savePath

    @staticmethod
    def createTable(doc, tableInfo):
        if tableInfo is None or len(tableInfo) <= 0:
            return
        table = doc.add_table(rows=1, cols=8, style="Light Grid Accent 1")
        table.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        table.style.font.size = Pt(10)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Result(m)'
        hdr_cells[1].text = 'mean'
        hdr_cells[2].text = 'std'
        hdr_cells[3].text = '68%'
        hdr_cells[4].text = '95%'
        hdr_cells[5].text = '99.7%'
        hdr_cells[6].text = 'max'
        hdr_cells[7].text = 'fixed'

        for name, cellInfo in tableInfo:
            cep_cell = table.add_row().cells
            cep_cell[0].text = name
            i = 0
            for row in cellInfo.iteritems():
                if i == 1:
                    cep_cell[1].text = str(f'%.3f' % row[1])
                elif i == 2:
                    cep_cell[2].text = str(f'%.3f' % row[1])
                elif i == 5:
                    cep_cell[3].text = str(f'%.3f' % row[1])
                elif i == 6:
                    cep_cell[4].text = str(f'%.3f' % row[1])
                elif i == 7:
                    cep_cell[5].text = str(f'%.3f' % row[1])
                elif i == 8:
                    cep_cell[6].text = str(f'%.3f' % row[1])
                elif i == 9:
                    cep_cell[7].text = str(f'%.2f%%' % row[1])
                i += 1

    def addPng(self, doc):
        for dirPath in os.listdir(self.ppPath):
            path_join = os.path.join(self.ppPath, dirPath)
            if os.path.isdir(path_join):
                for path in os.listdir(path_join):
                    resultPath = os.path.join(path_join, path)
                    if os.path.isdir(resultPath):
                        doc.add_paragraph('>>> %s:%s 图例 >>>' % (dirPath, path), style='Intense Quote')
                        for file in os.listdir(resultPath):
                            if file.endswith('png'):
                                doc.add_picture(os.path.join(resultPath, file))
                        doc.add_paragraph('<<< %s:%s 图例 <<<' % (dirPath, path), style='Intense Quote')

    def build(self, tableInfo=None):
        doc = Document()
        doc.add_heading('后处理版本测试', 0)
        doc.add_paragraph('测试报告生成时间：%s' % self.time)
        self.createTable(doc, tableInfo)
        self.addPng(doc)
        doc.save(self.savePath + '/ppResult.docx')


if __name__ == '__main__':
    pp = PPReporter(os.path.abspath('../../data/PP'), time="0001")
    pp.build()
